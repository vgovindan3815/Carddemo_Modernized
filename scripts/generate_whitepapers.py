from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(r"c:\Users\venkat.govindan\OneDrive - Accenture\personal\Projects\legacy-cobol-modernization-workspace")
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_subheading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# Business Perspective Whitepaper
business_doc = Document()
add_title(business_doc, "Legacy COBOL Modernization Whitepaper: Business Perspective")

business_doc.add_paragraph(
    "This whitepaper provides a business-focused view of the CardDemo modernization program, "
    "including online services and batch processing. It emphasizes outcomes, stakeholder "
    "value, risk posture, and execution approach for a pragmatic, phased modernization "
    "that preserves legacy behavior while enabling new digital capabilities."
)

add_heading(business_doc, "1. Executive Summary")
business_doc.add_paragraph(
    "The program modernizes a legacy COBOL CardDemo environment into a contemporary, "
    "service-oriented platform while preserving core business rules, auditability, and "
    "batch processing requirements. The result is faster change delivery, improved "
    "observability, and reduced operational risk without disrupting day-to-day operations."
)

add_heading(business_doc, "2. Business Objectives and Outcomes")
add_bullets(
    business_doc,
    [
        "Improve customer and agent experience through responsive, web-based interfaces and consistent workflows.",
        "Reduce time-to-market for regulatory and product updates by modularizing business logic into services.",
        "Preserve batch processing and file-based integrations required by downstream systems.",
        "Enhance auditability and traceability through structured APIs, centralized logging, and consistent error handling.",
        "Enable incremental modernization while minimizing disruption to existing operations.",
    ],
)

add_heading(business_doc, "3. Stakeholder Value")
add_subheading(business_doc, "Business and Operations")
add_bullets(
    business_doc,
    [
        "Greater transparency into account, card, transaction, billing, and reporting flows.",
        "Improved operational resilience with clearer runbooks and consistent error handling.",
        "Reduced training time due to modern UI patterns and guided actions.",
    ],
)
add_subheading(business_doc, "Technology and Delivery")
add_bullets(
    business_doc,
    [
        "Modern stack improves hiring, onboarding, and maintainability.",
        "Service boundaries and DTO contracts stabilize integration points.",
        "Automated migrations and seeds reduce environment drift.",
    ],
)

add_heading(business_doc, "4. Scope Overview")
add_bullets(
    business_doc,
    [
        "Online capabilities: account view/update, card view/update, transaction entry, billing, user administration, and reports.",
        "Batch capabilities: ingest and export of customer/account/card files, transaction posting, and operational backup outputs.",
        "Data store: SQLite for local persistence with migration support; design supports portability to managed databases.",
        "Security: authenticated sessions, role-based access, and controlled administrative paths.",
    ],
)

add_heading(business_doc, "5. Batch Processing in the Business Context")
business_doc.add_paragraph(
    "Batch processing remains central to the business due to daily settlement, reconciliation, "
    "and file exchange requirements. The modernization maintains batch behavior while introducing "
    "configurability for input/output paths and validation results. This protects downstream "
    "systems and preserves SLA commitments."
)
add_bullets(
    business_doc,
    [
        "File-driven ingestion for account, customer, and card datasets.",
        "Operational outputs support backups and audit trails.",
        "Optional keep-existing mode supports incremental loads without data loss.",
        "Clear return codes and logs enable operational triage and governance.",
    ],
)

add_heading(business_doc, "6. Risk Management")
add_bullets(
    business_doc,
    [
        "Phased rollout reduces risk by enabling feature validation before full cutover.",
        "Legacy parity checks ensure business rules are preserved.",
        "Batch controls and validation reduce file ingestion errors.",
        "Role-based access prevents unintended administrative changes.",
    ],
)

add_heading(business_doc, "7. Governance and Compliance")
add_bullets(
    business_doc,
    [
        "Structured error catalog and validation responses support auditability.",
        "API-level logging and correlation IDs improve traceability.",
        "Consistent data formats reduce compliance discrepancies.",
    ],
)

add_heading(business_doc, "8. Operating Model and Change Management")
add_bullets(
    business_doc,
    [
        "Modern UI reduces manual errors through validation and guided workflows.",
        "Batch runbooks document critical dependencies and rerun scenarios.",
        "Training can focus on new workflows rather than platform internals.",
    ],
)

add_heading(business_doc, "9. Value Realization Timeline")
add_bullets(
    business_doc,
    [
        "Short term: online modernization and stabilized batch ingestion.",
        "Mid term: expand reporting and administrative capabilities with faster releases.",
        "Long term: portability to managed databases and further integration automation.",
    ],
)

add_heading(business_doc, "10. Summary")
business_doc.add_paragraph(
    "This modernization delivers immediate operational improvements while protecting the "
    "batch-centric business workflows that underpin daily processing. The program positions "
    "the organization for faster change and lower risk across both online and batch domains."
)

business_doc.save(OUT_DIR / "CardDemo_Modernization_Business_Perspective.docx")

# Technical Whitepaper
tech_doc = Document()
add_title(tech_doc, "Legacy COBOL Modernization Whitepaper: Technical Perspective")
tech_doc.add_paragraph(
    "This whitepaper documents the technical architecture and implementation for the CardDemo modernization, "
    "including online services and batch processing. It focuses on system design, data flow, APIs, and "
    "operational considerations."
)

add_heading(tech_doc, "1. Architecture Overview")
add_bullets(
    tech_doc,
    [
        "Frontend: Angular single-page application with authenticated routes and shared components.",
        "Backend: Node.js + Express REST APIs implementing business logic and validation.",
        "Database: SQLite with migration and seed scripts for local development and repeatability.",
        "Batch processing: Node-based batch runner with file ingestion and export capabilities.",
    ],
)

add_heading(tech_doc, "2. Key Design Principles")
add_bullets(
    tech_doc,
    [
        "Legacy parity with incremental modernization.",
        "Strict DTO validation with consistent error payloads.",
        "Separation of concerns across UI, API, and data layers.",
        "Deterministic batch behavior with explicit return codes.",
    ],
)

add_heading(tech_doc, "3. Online API Surface")
add_bullets(
    tech_doc,
    [
        "Accounts: list, view detail, update account + customer.",
        "Cards: list, view detail, update.",
        "Transactions: list, view detail, create.",
        "Billing: post payments with card-on-file persistence.",
        "Users: admin create, update, delete.",
        "Reports: account summary and operational metrics.",
    ],
)

add_heading(tech_doc, "4. Data Model Highlights")
add_bullets(
    tech_doc,
    [
        "accounts: balances, limits, open/expiry/reissue dates, cycle credit/debit, group ID.",
        "customers: identity and contact data, address, primary holder indicator, FICO.",
        "cards: card data with status and expiration.",
        "card_xref: links between accounts, cards, and customers.",
        "transactions: posted activity supporting audit and reconciliation.",
    ],
)

add_heading(tech_doc, "5. Batch Processing Design")
tech_doc.add_paragraph(
    "Batch processing is implemented as a Node-based batch runner that reads legacy file formats, "
    "validates content, and performs deterministic upserts into the SQLite database. Outputs include "
    "backups and reconciliation-ready datasets."
)
add_bullets(
    tech_doc,
    [
        "Input files: ACCTFILE, CUSTFILE, CARDFILE (ASCII defaults supported).",
        "Processing modes: full replace or keep-existing for idempotent loads.",
        "Output: TRANBKP and export artifacts written to configured output paths.",
        "Validation: per-step metrics, errors, and return codes logged with timestamps.",
    ],
)

add_heading(tech_doc, "6. Security and Access Control")
add_bullets(
    tech_doc,
    [
        "Session-based authentication with secure cookies.",
        "Role-based access controls for admin operations.",
        "Input validation enforced with schema parsing to prevent malformed updates.",
    ],
)

add_heading(tech_doc, "7. Error Handling and Observability")
add_bullets(
    tech_doc,
    [
        "Consistent error catalog and response format.",
        "Correlation IDs in responses to improve traceability.",
        "Structured logging with clear context for batch jobs and APIs.",
    ],
)

add_heading(tech_doc, "8. Deployment and Environment")
add_bullets(
    tech_doc,
    [
        "Local run: combined dev mode for frontend and backend.",
        "Configurable API base URLs via environment settings.",
        "Database migrations and seed scripts ensure consistent environments.",
    ],
)

add_heading(tech_doc, "9. Operational Runbook Notes")
add_bullets(
    tech_doc,
    [
        "Batch runs accept input/output path overrides for file-based jobs.",
        "Return codes summarize validation success or failure per job.",
        "Logs capture totals, inserted/updated counts, and errors.",
    ],
)

add_heading(tech_doc, "10. Future Extensions")
add_bullets(
    tech_doc,
    [
        "Move SQLite to managed relational database with minimal schema changes.",
        "Introduce centralized job scheduling and monitoring.",
        "Add automated regression checks for batch parity.",
    ],
)

tech_doc.save(OUT_DIR / "CardDemo_Modernization_Technical_Perspective.docx")

# One-Pager
one_doc = Document()
add_title(one_doc, "Legacy COBOL Modernization: One-Pager")
one_doc.add_paragraph("A concise overview of the CardDemo modernization with online and batch scope.")

add_heading(one_doc, "What is being modernized")
add_bullets(
    one_doc,
    [
        "Online account, card, transaction, billing, and reporting workflows.",
        "Batch ingestion and export of customer, account, and card datasets.",
        "Legacy behavior preserved with updated UI and API integration.",
    ],
)

add_heading(one_doc, "Business Benefits")
add_bullets(
    one_doc,
    [
        "Faster change delivery and reduced operational risk.",
        "Improved visibility and auditability across online and batch processes.",
        "Lower maintenance cost through modern stack and automation.",
    ],
)

add_heading(one_doc, "Key Technical Elements")
add_bullets(
    one_doc,
    [
        "Angular frontend with authenticated routes and structured UI.",
        "Node.js/Express REST APIs with validation and consistent errors.",
        "SQLite database with migrations and seed scripts.",
        "Batch runner for file ingestion, validation, and outputs.",
    ],
)

add_heading(one_doc, "Batch Processing Highlights")
add_bullets(
    one_doc,
    [
        "File-driven ACCTFILE, CUSTFILE, CARDFILE ingestion.",
        "Optional keep-existing mode for safe, incremental loads.",
        "TRANBKP backup output and operational logging.",
    ],
)

add_heading(one_doc, "Next Steps")
add_bullets(
    one_doc,
    [
        "Finalize user training materials and operational runbooks.",
        "Validate batch parity using legacy comparison datasets.",
        "Plan database portability for enterprise deployment.",
    ],
)

one_doc.save(OUT_DIR / "CardDemo_Modernization_One_Pager.docx")

print(OUT_DIR / "CardDemo_Modernization_Business_Perspective.docx")
print(OUT_DIR / "CardDemo_Modernization_Technical_Perspective.docx")
print(OUT_DIR / "CardDemo_Modernization_One_Pager.docx")
